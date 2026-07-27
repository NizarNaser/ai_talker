"""
ملف واجهات برمجة التطبيقات (Views).
يحتوي على:
1. جلب وإنشاء الترجمات.
2. جلب وإضافة التعليقات.
3. الإعجاب بالموقع.
4. تسجيل الدخول باستخدام Google.
"""
from rest_framework import viewsets, views, status
from rest_framework.response import Response
from rest_framework.permissions import IsAuthenticated, AllowAny
from .models import User, Translation, SiteLike, Comment
from .serializers import TranslationSerializer, CommentSerializer
import logging

class TranslationViewSet(viewsets.ModelViewSet):
    """واجهة لإدارة الترجمات الخاصة بالمستخدم"""
    serializer_class = TranslationSerializer
    permission_classes = [IsAuthenticated]

    def get_queryset(self):
        return Translation.objects.filter(user=self.request.user)

    def perform_create(self, serializer):
        serializer.save(user=self.request.user)


class CommentViewSet(viewsets.ModelViewSet):
    """واجهة لعرض وإضافة التعليقات"""
    serializer_class = CommentSerializer
    
    def get_permissions(self):
        if self.request.method in ['POST', 'PUT', 'PATCH', 'DELETE']:
            return [IsAuthenticated()]
        return [AllowAny()]

    def get_queryset(self):
        return Comment.objects.filter(is_approved=True)

    def perform_create(self, serializer):
        serializer.save(user=self.request.user)


class SiteLikeView(views.APIView):
    """واجهة للتعامل مع الإعجابات بالموقع"""
    permission_classes = [AllowAny]

    def get(self, request):
        like_obj, created = SiteLike.objects.get_or_create(id=1)
        return Response({'total_likes': like_obj.total_likes})

    def post(self, request):
        like_obj, created = SiteLike.objects.get_or_create(id=1)
        like_obj.total_likes += 1
        like_obj.save()
        return Response({'total_likes': like_obj.total_likes, 'message': 'تم إضافة الإعجاب بنجاح.'})


class GoogleLoginView(views.APIView):
    """واجهة وهمية لتسجيل الدخول عبر جوجل (سيتم تطويرها لاحقاً لربط OAuth)"""
    permission_classes = [AllowAny]

    def post(self, request):
        token = request.data.get("token")
        # في التطبيق الحقيقي، سنقوم بالتحقق من token باستخدام مكتبة google.oauth2.id_token
        # إذا كان صحيحاً، نقوم بتسجيل المستخدم وإصدار JWT
        return Response({"message": "تم استلام الطلب", "token": token}, status=status.HTTP_200_OK)

class RegisterView(views.APIView):
    """واجهة لتسجيل مستخدم جديد"""
    permission_classes = [AllowAny]

    def post(self, request):
        email = request.data.get('email')
        password = request.data.get('password')
        if not email or not password:
            return Response({'error': 'البريد الإلكتروني وكلمة المرور مطلوبان'}, status=status.HTTP_400_BAD_REQUEST)
        
        # We will use email as username
        if User.objects.filter(username=email).exists():
            # If user exists, we don't return error, we assume they want to login later.
            return Response({'error': 'المستخدم موجود مسبقاً'}, status=status.HTTP_400_BAD_REQUEST)
            
        User.objects.create_user(username=email, email=email, password=password)
        return Response({'message': 'تم إنشاء الحساب بنجاح'}, status=status.HTTP_201_CREATED)

class ContactView(views.APIView):
    """واجهة لإرسال رسائل تواصل معنا عبر البريد الإلكتروني"""
    permission_classes = [AllowAny]

    def post(self, request):
        from django.core.mail import send_mail
        email = request.data.get('email')
        message = request.data.get('message')
        
        if not email or not message:
            return Response({'error': 'البريد والرسالة مطلوبان'}, status=status.HTTP_400_BAD_REQUEST)
        
        try:
            send_mail(
                subject=f"رسالة تواصل من: {email}",
                message=message,
                from_email=email,
                recipient_list=['admin@aitalker.local'],
                fail_silently=False,
            )
            return Response({'message': 'تم إرسال رسالتك بنجاح. شكراً لتواصلك معنا!'}, status=status.HTTP_200_OK)
        except Exception:
            return Response({'error': 'حدث خطأ أثناء إرسال الرسالة.'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

class FileUploadTranslateView(views.APIView):
    """واجهة لرفع ملف أو صورة، استخراج النص، وترجمته مع الحفاظ على التنسيق للمستندات"""
    permission_classes = [AllowAny]

    def post(self, request):
        logger = logging.getLogger(__name__)
        logger.info('📥 Received file upload request from %s', request.META.get('REMOTE_ADDR'))
        file_obj = request.FILES.get('file')
        target_lang = request.data.get('target_lang', 'ar')
        source_lang = request.data.get('source_lang', 'auto')
        
        if not file_obj:
            return Response({'error': 'لم يتم العثور على ملف في الطلب.'}, status=status.HTTP_400_BAD_REQUEST)

        # Preliminary checks
        logger.info('📎 File: %s, Size: %s bytes', file_obj.name, file_obj.size)
        max_file_size = 5 * 1024 * 1024  # 5 MiB
        if file_obj.size > max_file_size:
            logger.warning('File size %s exceeds limit %s', file_obj.size, max_file_size)
            return Response({'error': 'حجم الملف كبير جداً. أقصى حجم مسموح هو 8 MiB.'}, status=status.HTTP_400_BAD_REQUEST)

        extracted_text = ""
        translated_text = ""
        translated_file_base64 = None
        translated_file_format = None
        file_name = file_obj.name.lower()
        file_bytes = file_obj.read()
        # Track processing time to avoid long hangs
        import time
        processing_start = time.time()
        max_processing_seconds = 30

        def check_timeout():
            if time.time() - processing_start > max_processing_seconds:
                raise TimeoutError(f'Processing time exceeded {max_processing_seconds} seconds')


        try:
            # Setup Translator
            from deep_translator import GoogleTranslator
            
            def map_lang(lang):
                if not lang: return 'auto'
                lang_lower = str(lang).strip().lower()
                if lang_lower in ['zh', 'zh-cn', 'chinese', 'chinese (simplified)', 'zh-hans']:
                    return 'zh-CN'
                if lang_lower in ['zh-tw', 'chinese (traditional)', 'zh-hant']:
                    return 'zh-TW'
                if lang_lower.startswith('ar-'):
                    return 'ar'
                return str(lang).strip()

            safe_source = map_lang(source_lang)
            safe_target = map_lang(target_lang)
            if safe_source == 'zh': safe_source = 'zh-CN'
            if safe_target == 'zh': safe_target = 'zh-CN'
            
            translator = GoogleTranslator(source=safe_source, target=safe_target)

            def process_docx(bytes_data):
                import docx
                import io
                import base64
                doc = docx.Document(io.BytesIO(bytes_data))
                full_original = []
                full_translated = []
                
                # Helper to translate safely
                def translate_text(text):
                    if not text.strip(): return text
                    try:
                        res = translator.translate(text)
                        return res if res else text
                    except Exception:
                        return text

                for p in doc.paragraphs:
                    if p.text.strip():
                        original = p.text
                        full_original.append(original)
                        trans = translate_text(original)
                        full_translated.append(trans)
                        p.text = trans
                        
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for p in cell.paragraphs:
                                if p.text.strip():
                                    original = p.text
                                    full_original.append(original)
                                    trans = translate_text(original)
                                    full_translated.append(trans)
                                    p.text = trans
                                    
                output = io.BytesIO()
                doc.save(output)
                return "\n".join(full_original), "\n".join(full_translated), base64.b64encode(output.getvalue()).decode('utf-8')

            # Process Document Types
            if file_name.endswith('.docx'):
                try:
                    extracted_text, translated_text, translated_file_base64 = process_docx(file_bytes)
                    translated_file_format = 'docx'
                    check_timeout()
                except Exception as e:
                    return Response({'error': f'حدث خطأ أثناء معالجة ملف Word: {str(e)}'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

            elif file_name.endswith('.pdf'):
                try:
                    import tempfile
                    import os
                    from pdf2docx import Converter
                    
                    with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as pdf_file:
                        pdf_file.write(file_bytes)
                        pdf_path = pdf_file.name
                        
                    docx_path = pdf_path + '.docx'
                    try:
                        cv = Converter(pdf_path)
                        cv.convert(docx_path, start=0, end=None)
                        cv.close()
                        
                        with open(docx_path, 'rb') as f:
                            docx_bytes = f.read()
                            
                        extracted_text, translated_text, translated_file_base64 = process_docx(docx_bytes)
                        translated_file_format = 'docx' # Returns a docx
                        check_timeout()
                    finally:
                        if os.path.exists(pdf_path): os.remove(pdf_path)
                        if os.path.exists(docx_path): os.remove(docx_path)
                except ImportError:
                    return Response({'error': 'مكتبات تحويل PDF غير مثبتة في الخادم (pdf2docx).'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
                except Exception as e:
                    return Response({'error': f'حدث خطأ أثناء معالجة ملف PDF: {str(e)}'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
                
            elif file_name.endswith(('.png', '.jpg', '.jpeg', '.webp')):
                try:
                    import pytesseract
                    from PIL import Image
                    import io
                    import shutil as _shutil

                    # تحديد مسار tesseract تلقائياً (Render يثبته في /usr/bin)
                    _tess_path = _shutil.which('tesseract')
                    if not _tess_path:
                        import os as _os
                        for _candidate in ['/usr/bin/tesseract', '/usr/local/bin/tesseract']:
                            if _os.path.isfile(_candidate):
                                _tess_path = _candidate
                                break
                    if _tess_path:
                        pytesseract.pytesseract.tesseract_cmd = _tess_path
                        logger.info('✅ Tesseract found at: %s', _tess_path)
                    else:
                        logger.error('❌ Tesseract binary not found')
                        return Response(
                            {'error': 'محرك التعرف على النصوص (Tesseract) غير مثبت على الخادم.'},
                            status=status.HTTP_503_SERVICE_UNAVAILABLE
                        )

                    image = Image.open(io.BytesIO(file_bytes)).convert('RGB')
                    file_bytes = None  # تحرير الذاكرة فوراً

                    # تصغير الصورة لتوفير الذاكرة على Render
                    max_dim = 1600
                    w, h = image.size
                    if w > max_dim or h > max_dim:
                        ratio = min(max_dim / w, max_dim / h)
                        image = image.resize((int(w * ratio), int(h * ratio)), Image.LANCZOS)
                        logger.info('🖼️ Resized image from %dx%d to %dx%d', w, h, int(w*ratio), int(h*ratio))

                    # تحويل رمز اللغة إلى صيغة Tesseract
                    tesseract_lang_map = {
                        'ar': 'ara', 'en': 'eng', 'fr': 'fra', 'es': 'spa',
                        'de': 'deu', 'ru': 'rus', 'zh-CN': 'chi_sim', 'zh-TW': 'chi_tra',
                        'ja': 'jpn', 'ko': 'kor', 'it': 'ita', 'pt': 'por',
                        'nl': 'nld', 'sv': 'swe', 'tr': 'tur', 'hi': 'hin',
                        'auto': 'ara+eng',
                    }
                    tess_lang = tesseract_lang_map.get(safe_source, 'ara+eng')

                    extracted_text = pytesseract.image_to_string(image, lang=tess_lang).strip()
                    del image

                    if not extracted_text:
                        return Response({'error': 'لم يتم العثور على نص في الصورة. تأكد أن الصورة واضحة وتحتوي على نص مقروء.'}, status=status.HTTP_400_BAD_REQUEST)

                    # ترجمة بالدفعات
                    chunk_size = 4500
                    translated_chunks = []
                    for i in range(0, len(extracted_text), chunk_size):
                        chunk = extracted_text[i:i+chunk_size]
                        res = translator.translate(chunk)
                        if res: translated_chunks.append(res)

                    translated_text = ' '.join(translated_chunks)
                    check_timeout()

                except ImportError:
                    return Response({'error': 'مكتبة pytesseract غير مثبتة في الخادم.'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
                except Exception as e:
                    logger.error('OCR error: %s', str(e))
                    return Response({'error': f'خطأ في معالجة الصورة: {str(e)}'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
            
            else:
                # For unsupported file types: save the file to the configured storage
                # (Cloudinary or local MEDIA storage) and return a download URL so the
                # frontend can still upload arbitrary files even if we don't process them.
                try:
                    from django.core.files.base import ContentFile
                    from django.core.files.storage import default_storage
                    import os
                    import uuid

                    uploads_dir = 'uploads'
                    base_name = file_obj.name
                    save_path = os.path.join(uploads_dir, base_name)
                    # ensure unique path
                    if default_storage.exists(save_path):
                        name, ext = os.path.splitext(base_name)
                        save_path = os.path.join(uploads_dir, f"{name}_{uuid.uuid4().hex}{ext}")

                    default_storage.save(save_path, ContentFile(file_bytes))
                    file_url = default_storage.url(save_path)

                    return Response({
                        'message': 'تم رفع الملف بنجاح (نوع غير مدعوم للمعالجة).',
                        'file_url': file_url,
                    }, status=status.HTTP_200_OK)
                except Exception as e:
                    return Response({'error': f'فشل رفع الملف: {str(e)}'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

            return Response({
                'original_text': extracted_text,
                'translated_text': translated_text,
                'translated_file_base64': translated_file_base64,
                'translated_file_format': translated_file_format,
                'message': 'تم استخراج وترجمة النص بنجاح'
            }, status=status.HTTP_200_OK)

        except TimeoutError as te:
            logger.warning('Processing timeout: %s', str(te))
            return Response({'error': 'انتهت مهلة معالجة الملف. يرجى رفع ملف أصغر أو تجربة مرة أخرى.'}, status=status.HTTP_504_GATEWAY_TIMEOUT)
        except Exception as e:
            return Response({'error': f'حدث خطأ غير متوقع: {str(e)}'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


class HealthCheckView(views.APIView):
    """Simple health check endpoint to verify HTTP server reachability."""
    permission_classes = [AllowAny]

    def get(self, request):
        return Response({
            'status': 'ok',
            'message': 'Backend HTTP server is reachable',
            'ws_test_path': '/ws/translate/'
        })
